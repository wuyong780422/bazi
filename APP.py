# ===================== 最终版：顶部标题栏（图标完整+点击响应+零报错） =====================
import streamlit as st

# 隐藏Streamlit默认顶部栏
st.markdown("""
<style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    /* 自定义固定顶部标题栏 */
    .fixed-header {
        position: fixed;
        top: 0;
        left: 0;
        width: 100%;
        background: white;
        padding: 10px 20px;
        border-bottom: 1px solid #eee;
        z-index: 9999;
        display: flex;
        align-items: center;
        justify-content: space-between;
        height: 60px;
    }
    .header-icon {
        width: 36px;
        height: 36px;
        display: flex;
        align-items: center;
        justify-content: center;
        cursor: pointer;
    }
    .header-title {
        position: absolute;
        left: 50%;
        transform: translateX(-50%);
        font-size: 22px;
        font-weight: bold;
        color: #333;
    }
    .header-menu {
        font-size: 24px;
        color: #333;
    }
    /* 联系我们弹窗 */
    .wechat-modal {
        position: fixed;
        top: 0;
        left: 0;
        width: 100%;
        height: 100%;
        background: rgba(0,0,0,0.7);
        z-index: 99999;
        display: none;
        align-items: center;
        justify-content: center;
    }
    .modal-content {
        background: white;
        padding: 30px;
        border-radius: 16px;
        width: 300px;
        text-align: center;
    }
    /* 给页面主体加顶部内边距，避免被标题栏挡住 */
    .block-container {
        padding-top: 80px !important;
    }
</style>
""", unsafe_allow_html=True)

# 顶部标题栏（图标居中+完整显示）
st.markdown("""
<div class="fixed-header">
    <div class="header-icon" id="wechat_icon">
        <svg width="28" height="28" viewBox="0 0 24 24" fill="#07C160">
            <path d="M8.691 2.188C3.891 2.188 0 5.476 0 9.53c0 2.212 1.17 4.203 3.002 5.55a.59.59 0 0 1 .213.665l-.39 1.48c-.019.07-.048.141-.048.213 0 .163.13.295.29.295a.326.326 0 0 0 .167-.054l1.903-1.114a.864.864 0 0 1 .717-.098 10.16 10.16 0 0 0 2.837.403c.276 0 .543-.027.811-.05-.857-2.578.157-4.972 1.932-6.446 1.703-1.415 3.882-1.98 5.853-1.838-.576-3.583-4.196-6.348-8.596-6.348zM5.785 5.991c.642 0 1.162.529 1.162 1.18a1.17 1.17 0 0 1-1.162 1.178A1.17 1.17 0 0 1 4.623 7.17c0-.651.52-1.18 1.162-1.18zm5.813 0c.642 0 1.162.529 1.162 1.18a1.17 1.17 0 0 1-1.162 1.178 1.17 1.17 0 0 1-1.162-1.178c0-.651.52-1.18 1.162-1.18zm5.34 2.867c-1.797-.052-3.746.512-5.28 1.786-1.72 1.428-2.687 3.72-1.78 6.22.942 2.453 3.666 4.229 6.884 4.229.826 0 1.622-.12 2.361-.336a.722.722 0 0 1 .598.082l1.584.926a.272 0 0 0 .14.047c.134 0 .24-.111.24-.247 0-.06-.023-.12-.038-.177l-.327-1.233a.582.582 0 0 1-.023-.156.49.49 0 0 1 .201-.398C23.024 18.48 24 16.82 24 14.98c0-3.21-2.931-5.837-6.656-6.088V8.89c-.135-.01-.27-.027-.407-.03zm-2.53 3.274c.535 0 .969.44.969.982a.976.976 0 0 1-.969.983.976.976 0 0 1-.969-.983c0-.542.434-.982.97-.982zm4.844 0c.535 0 .969.44.969.982a.976.976 0 0 1-.969.983.976.976 0 0 1-.969-.983c0-.542.434-.982.969-.982z"/>
        </svg>
    </div>
    <span class="header-title">真命盘专业版</span>
    <span class="header-menu">⋯</span>
</div>

<!-- 联系我们弹窗 -->
<div id="wechat_modal" class="wechat-modal">
    <div class="modal-content">
        <h3>📞 联系我们</h3>
        <img src="wechat.png" width="220" style="margin:10px 0;">
        <p>扫描二维码添加微信</p>
        <button onclick="document.getElementById('wechat_modal').style.display='none'"
        style="background:#07C160; color:white; border:none; padding:10px 25px; border-radius:10px; cursor:pointer;">
        关闭
        </button>
    </div>
</div>

<!-- JS 点击事件绑定 -->
<script>
document.getElementById('wechat_icon').addEventListener('click', function(){
    document.getElementById('wechat_modal').style.display = 'flex';
});
</script>
""", unsafe_allow_html=True)
# ==========================================================
# 真命盘专业版 —— 固定顶部标题+全功能原版+手机端历法一行优化版
#  运行命令>> streamlit run APP.py
# ==========================================================
import os
import sys
import requests
from datetime import  timedelta
import sqlite3
# ===================== 资源路径（和APP0完全一样） =====================
def resource_path(relative_path: str) -> str:
    try:
        base_path = sys._MEIPASS
    except AttributeError:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)
# ===================== 命理常量（完全不变） =====================
SHICHEN_DETAIL = [
    "子时 23:00-01:00", "丑时 01:00-03:00", "寅时 03:00-05:00", "卯时 05:00-07:00",
    "辰时 07:00-09:00", "巳时 09:00-11:00", "午时 11:00-13:00", "未时 13:00-15:00",
    "申时 15:00-17:00", "酉时 17:00-19:00", "戌时 19:00-21:00", "亥时 21:00-23:00"
]
SHICHEN_TIME = [s.split(" ")[0] for s in SHICHEN_DETAIL]
TIANGAN_WUXING = {"甲": "木", "乙": "木", "丙": "火", "丁": "火", "戊": "土", "己": "土", "庚": "金", "辛": "金", "壬": "水", "癸": "水"}
YUELING_WUXING = {"寅": "木", "卯": "木", "辰": "土", "巳": "火", "午": "火", "未": "土", "申": "金", "酉": "金", "戌": "土", "亥": "水", "子": "水", "丑": "土"}
NAYIN_TABLE = {
    "甲子": "海中金", "乙丑": "海中金", "丙寅": "炉中火", "丁卯": "炉中火", "戊辰": "大林木", "己巳": "大林木",
    "庚午": "路旁土", "辛未": "路旁土", "壬申": "剑锋金", "癸酉": "剑锋金", "甲戌": "山头火", "乙亥": "山头火",
    "丙子": "涧下水", "丁丑": "涧下水", "戊寅": "城头土", "己卯": "城头土", "庚辰": "白蜡金", "辛巳": "白蜡金",
    "壬午": "杨柳木", "癸未": "杨柳木", "甲申": "泉中水", "乙酉": "泉中水", "丙戌": "屋上土", "丁亥": "屋上土",
    "戊子": "霹雳火", "己丑": "霹雳火", "庚寅": "松柏木", "辛卯": "松柏木", "壬辰": "长流水", "癸巳": "长流水",
    "甲午": "沙中金", "乙未": "沙中金", "丙申": "山下火", "丁酉": "山下火", "戊戌": "平地木", "己亥": "平地木",
    "庚子": "壁上土", "辛丑": "壁上土", "壬寅": "金箔金", "癸卯": "金箔金", "甲辰": "佛灯火", "乙巳": "佛灯火",
    "丙午": "天河水", "丁未": "天河水", "戊申": "大驿土", "己酉": "大驿土", "庚戌": "钗钏金", "辛亥": "钗钏金",
    "壬子": "桑柘木", "癸丑": "桑柘木", "甲寅": "大溪水", "乙卯": "大溪水", "丙辰": "沙中土", "丁巳": "沙中土",
    "戊午": "天上火", "己未": "天上火", "庚申": "石榴木", "辛酉": "石榴木", "壬戌": "大海水", "癸亥": "大海水"
}
SHI_ZHI = ["子","丑","寅","卯","辰","巳","午","未","申","酉","戌","亥"]
JX_MAP = {"建":"次吉","除":"大吉","满":"平","平":"平","定":"吉","执":"平","破":"大凶","危":"平","成":"大吉","收":"吉","开":"大吉","闭":"平"}

# ===================== 民俗风水专用常量 =====================
SHENGXIAO_WUXING = {
    "鼠": "水", "牛": "土", "虎": "木", "兔": "木", "龙": "土", "蛇": "火",
    "马": "火", "羊": "土", "猴": "金", "鸡": "金", "狗": "土", "猪": "水"
}
YONGSHEN_FANGWEI = {"木": "东", "火": "南", "土": "中", "金": "西", "水": "北"}
LOU_CENG_WUXING = {"水": [1, 6], "火": [2, 7], "木": [3, 8], "金": [4, 9], "土": [5, 10]}
BAZHAI_JIXIONG = ["生气", "天医", "延年", "伏位", "绝命", "五鬼", "六煞", "祸害"]

# ===================== 【极简稳定版】导出模块（仅保留Word，无PDF依赖问题） =====================
import io
from datetime import datetime
# import streamlit as st

# ---------------------- Word导出（无兼容问题，100%可用） ----------------------
DOCX_AVAILABLE = False
Document = None
try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    pass
def generate_word_doc(bazi_data, gender, ai_content, fengshui_content=""):
    if not DOCX_AVAILABLE:
        return None
    doc = Document()
    # 标题居中（兼容所有python-docx版本）
    title = doc.add_heading("八字命理综合测算报告", 0)
    title_paragraph = doc.paragraphs[0]
    title_paragraph.alignment = 1  # 1=居中，无需导入常量

    doc.add_heading("一、基础信息", level=1)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"性别：{gender}")
    doc.add_paragraph(f"公历：{bazi_data['公历']}")
    doc.add_paragraph(f"农历：{bazi_data['农历']}")
    doc.add_paragraph(f"生肖：{bazi_data['生肖']}")
    doc.add_paragraph(f"完整八字：{bazi_data['八字_str']}")
    doc.add_paragraph(f"日主：{bazi_data['日干']}({bazi_data['日干五行']})")

    doc.add_heading("二、AI深度解读", level=1)
    doc.add_paragraph(ai_content)
    if fengshui_content:
        doc.add_heading("三、民俗风水布局", level=1)
        doc.add_paragraph(fengshui_content)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ===================== 天干五行辅助函数（补回bazi_result['日干五行']专用） =====================
def get_gan_wuxing(gan: str) -> str:
    """
    根据天干返回传统命理中的五行属性
    规则：甲乙属木，丙丁属火，戊己属土，庚辛属金，壬癸属水
    """
    if gan in ["甲", "乙"]:
        return "木"
    elif gan in ["丙", "丁"]:
        return "火"
    elif gan in ["戊", "己"]:
        return "土"
    elif gan in ["庚", "辛"]:
        return "金"
    elif gan in ["壬", "癸"]:
        return "水"
    else:
        return ""
# ===================== AI解读专用配置（和PC端完全一致） =====================
DEEPSEEK_CONFIG = {
    "api_key": "sk-d3b976c41272460eab55726001606b15",
    "api_url": "https://api.deepseek.com/v1/chat/completions",
    "model": "deepseek-chat",
    "timeout": 70,
    "temperature": 0.7,
    "max_tokens": 2800
}
# ===================== 数据库函数（100% 复制 APP0 正确版） =====================
def query_db_ganzhi(solar_date_str: str) -> tuple[str, str, str]:
    try:
        conn = sqlite3.connect(resource_path("bazi_calendar.db"), timeout=10)
        cursor = conn.cursor()
        cursor.execute('SELECT 年, 月建, 纳音 FROM calendar WHERE 国历 = ? LIMIT 1', (solar_date_str,))
        result = cursor.fetchone()
        conn.close()
        return result if result else ("甲子", "甲子", "甲子")
    except Exception:
        return "甲子", "甲子", "甲子"

def solar_to_lunar_from_db(solar_date_str: str) -> dict[str, str]:
    try:
        conn = sqlite3.connect(resource_path("bazi_calendar.db"), timeout=10)
        cursor = conn.cursor()
        cursor.execute('SELECT 农历年,农历月,农历日,生肖 FROM calendar WHERE 国历 = ? LIMIT 1', (solar_date_str,))
        res = cursor.fetchone()
        conn.close()
        if res:
            return {"农历完整信息": f"{res[0]}年{res[1]}月{res[2]}日", "生肖": res[3]}
        else:
            return {"农历完整信息": "2026年三月初十", "生肖": "马"}
    except Exception:
        return {"农历完整信息": "2026年三月初十", "生肖": "马"}

def lunar_to_solar_from_db(lunar_year: int, lunar_month: int, lunar_day: int, is_leap: int) -> str:
    try:
        conn = sqlite3.connect(resource_path("bazi_calendar.db"), timeout=10)
        cursor = conn.cursor()
        cursor.execute('''SELECT 国历 FROM calendar WHERE 农历年=? AND 农历月=? AND 农历日=? AND 闰月=? LIMIT 1''',
                       (lunar_year, lunar_month, lunar_day, is_leap))
        res = cursor.fetchone()
        conn.close()
        return res[0] if res else "1990-01-01"
    except Exception:
        return "1990-01-01"

def calculate_shichen_ganzhi(ri_gan: str, shichen: str) -> str:
    shichen_zhi_map = {"子时": "子", "丑时": "丑", "寅时": "寅", "卯时": "卯", "辰时": "辰", "巳时": "巳", "午时": "午",
                       "未时": "未", "申时": "申", "酉时": "酉", "戌时": "戌", "亥时": "亥"}
    rigan_zi_shigan_map = {"甲": "甲", "己": "甲", "乙": "丙", "庚": "丙", "丙": "戊", "辛": "戊", "丁": "庚",
                           "壬": "庚", "戊": "壬", "癸": "壬"}
    gan_list = ["甲", "乙", "丙", "丁", "戊", "己", "庚", "辛", "壬", "癸"]
    shichen_index = list(shichen_zhi_map.keys()).index(shichen)
    zi_gan_index = gan_list.index(rigan_zi_shigan_map[ri_gan])
    return gan_list[(zi_gan_index + shichen_index) % 10] + shichen_zhi_map[shichen]

# ===================== 八字核心计算（完全不变） =====================
class BaziCalculator:
    @staticmethod
    def generate_bazi(target_date: str, shichen: str) -> dict:
        nian, yue, ri = query_db_ganzhi(target_date)
        ri_gan = ri[0]
        shi = calculate_shichen_ganzhi(ri_gan, shichen)
        lunar = solar_to_lunar_from_db(target_date)
        wuxing = {"金": 0, "木": 0, "水": 0, "火": 0, "土": 0}
        for g in [nian[0], yue[0], ri[0], shi[0]]:
            wuxing[TIANGAN_WUXING[g]] += 1
        for z in [nian[1], yue[1], ri[1], shi[1]]:
            wuxing[YUELING_WUXING[z]] += 1
        return {
            "公历": target_date, "农历": lunar["农历完整信息"], "生肖": lunar["生肖"],
            "时辰": shichen, "八字": [nian, yue, ri, shi], "八字_str": f"{nian} {yue} {ri} {shi}",
            "日干": ri_gan,
            # 👇 【唯一新增：补全日干五行，和电脑版完全一致】
            "日干五行": get_gan_wuxing(ri_gan),
            "五行": wuxing,
            "纳音": [NAYIN_TABLE.get(nian, ""), NAYIN_TABLE.get(yue, ""), NAYIN_TABLE.get(ri, ""),
                     NAYIN_TABLE.get(shi, "")]
        }

    @staticmethod
    def get_current_bazi() -> dict:
        now = datetime.now()
        current_date = now.strftime("%Y-%m-%d")
        current_hour = now.hour
        shichen_map = {23: "子时", 0: "子时", 1: "丑时", 2: "丑时", 3: "寅时", 4: "寅时", 5: "卯时", 6: "卯时",
                       7: "辰时", 8: "辰时", 9: "巳时", 10: "巳时", 11: "午时", 12: "午时", 13: "未时", 14: "未时",
                       15: "申时", 16: "申时", 17: "酉时", 18: "酉时", 19: "戌时", 20: "戌时", 21: "亥时", 22: "亥时"}
        current_shichen = shichen_map.get(current_hour, "亥时")
        return BaziCalculator.generate_bazi(current_date, current_shichen)

# ===================== 页面样式（完全保留你的原版，仅优化手机端历法布局） =====================
st.set_page_config(page_title="真命盘专业版", page_icon="☯️", layout="centered")
page_bg = """
<style>
body {margin: 0;padding: 0;background-color:#E5E5E5;}
.stApp {background-color:#E5E5E5;padding: 0;}
div.block-container {padding-top: 50px !important;padding-left:20px !important;padding-right:20px !important;padding-bottom:60px !important;}
div.stContainer {background:#F0F0F0;border-radius:12px;padding:15px;}
div.stTextInput>div>div {border-radius:8px; background:#FFF;}
div.stSelectbox>div>div {border-radius:8px; background:#FFF;}
div.stDateInput>div>div {border-radius:8px; background:#FFF;}
div.stNumberInput>div>div {border-radius:8px; background:#FFF;}
/* 关键修复：只让radio自适应，不影响按钮 */
div.stRadio>div {display:flex;gap:8px;justify-content:center;flex-wrap:wrap !important;max-width:100%;}
div.stRadio label {background:#FFF;border-radius:20px;padding:8px 16px;border:1px solid #EEE;font-size:14px;white-space:nowrap;box-sizing:border-box;}
div.stRadio [role="radio"]:checked + label {background:#D4AF37;color:#FFF;border-color:#D4AF37;}
div.stButton>button {background-color:#222222;color:#D4AF37;border-radius:30px;height:68px;font-size:18px;font-weight:bold;width:100%;}
.footer-nav {position:fixed;bottom:0;left:0;right:0;background:#FFF;display:flex;justify-content:space-around;padding:10px 0;border-top:1px solid #EEE;z-index:100;}
.nav-item {text-align:center;font-size:12px;color:#666;}
.nav-item.active {color:#9370DB;}
</style>
"""
st.markdown(page_bg, unsafe_allow_html=True)
# 页面状态控制：排盘=主排盘页，吉日=吉日页,风水=风水页，解读=解读页
if "bottom_nav_active" not in st.session_state:
    st.session_state.bottom_nav_active = "排盘"
# ===================== 界面与功能（仅优化性别/历法布局，和性别一样一行显示） =====================
with st.container(border=True):
    col_name_label, col_name_input = st.columns([1, 4])
    with col_name_label: st.markdown("**姓名**")
    with col_name_input: name = st.text_input("", placeholder="请输入姓名", label_visibility="collapsed",key="name" )

    # 优化点：性别和历法放在同一行两列，和性别一样一行显示
    col_gender, col_cal = st.columns(2)
    with col_gender:
        st.markdown("**性别**")
        gender = st.radio("", ["先生", "女士"], horizontal=True, label_visibility="collapsed")
        # 👇 修复核心：把性别存入session_state，风水页面才能读到
        st.session_state.gender = gender
    with col_cal:
        st.markdown("**历法**")
        calendar_type = st.radio("", ["公历", "农历"], horizontal=True, label_visibility="collapsed")

    if calendar_type == "公历":
        st.markdown("**出生时间（必填）**")
        birth_date = st.date_input("", datetime(2000, 1, 1), min_value=datetime(1900, 1, 1), max_value=datetime(2100, 12, 31), label_visibility="collapsed")
        date_str = birth_date.strftime("%Y-%m-%d")
    else:
        col_lun_year, col_lun_month, col_lun_day = st.columns(3)
        with col_lun_year:
            st.markdown("**农历年**")
            lunar_year_input = st.number_input("", 1900, 2100, 2000, label_visibility="collapsed")
        with col_lun_month:
            st.markdown("**农历月**")
            lunar_month_input = st.number_input("", 1, 12, 6, label_visibility="collapsed")
        with col_lun_day:
            st.markdown("**农历日**")
            lunar_day_input = st.number_input("", 1, 30, 15, label_visibility="collapsed")
        st.markdown("**平月/闰月**")
        leap_option = st.radio("", ["平月", "闰月"], horizontal=True, label_visibility="collapsed")
        is_leap_input = 1 if leap_option == "闰月" else 0
        date_str = lunar_to_solar_from_db(lunar_year_input, lunar_month_input, lunar_day_input, is_leap_input)

    st.markdown("**出生地区**")
    birth_area = st.selectbox("", ["北京", "成都", "上海", "广州", "深圳"], index=0, label_visibility="collapsed")
    true_sun_time = "1990-01-01 00:00"
    lat, lon = "北纬39.93", "东经116.42"

    st.markdown("**出生时辰**")

    # 分成两段，彻底避开 Streamlit 吞选项 bug
    shi_group = st.radio("", ["子-巳", "午-亥"], horizontal=True, label_visibility="collapsed")

    if shi_group == "子-巳":
        selected_shichen_detail = st.selectbox("", [
            "子时 23:00-01:00", "丑时 01:00-03:00", "寅时 03:00-05:00",
            "卯时 05:00-07:00", "辰时 07:00-09:00", "巳时 09:00-11:00"
        ], index=0, label_visibility="collapsed")
    else:
        selected_shichen_detail = st.selectbox("", [
            "午时 11:00-13:00", "未时 13:00-15:00", "申时 15:00-17:00",
            "酉时 17:00-19:00", "戌时 19:00-21:00", "亥时 21:00-23:00"
        ], index=5, label_visibility="collapsed")

    shichen_input = selected_shichen_detail.split(" ")[0]

    # 按钮布局优化版：1:1等宽 + 小间距，仅作用于这两个按钮
    col_btn1, col_btn2 = st.columns(2, gap="small")
    with col_btn1:
        if st.button("开始排盘", use_container_width=True):
            st.session_state.bazi_result = BaziCalculator.generate_bazi(date_str, shichen_input)
    with col_btn2:
        if st.button("即时排盘", use_container_width=True):
            st.session_state.bazi_result = BaziCalculator.get_current_bazi()

    col_info, col_save = st.columns([3, 1])
    with col_info:
        st.markdown(f"""<div style="color:#666;font-size:12px;">真太阳时：{true_sun_time}<br>地址经纬：{lat} {lon}</div>""", unsafe_allow_html=True)
    with col_save:
        save_toggle = st.toggle("保存", value=False)

# ========== 四柱显示：居中+字体大小自由控制 ==========
if "bazi_result" in st.session_state and st.session_state.bazi_result:
        r = st.session_state.bazi_result
        st.markdown("---")
        st.success("✅ 排盘完成")

        # 用HTML表格，完全自定义样式
        pillars_html = f"""
        <style>
        .custom-table {{
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
        }}
        .custom-table th, .custom-table td {{
            border: 1px solid #ddd;
            padding: 12px 8px;
            text-align: center;
        }}
        .custom-table th {{
            font-size: 14px; /* 表头文字大小 */
            color: #666;
            background-color: #f8f8f8;
        }}
        .custom-table td {{
            font-size: 28px; /* ← 八字文字大小，改这里！ */
            font-weight: bold;
            color: #333;
        }}
        @media (max-width: 600px) {{
            .custom-table th {{ font-size: 12px; }}
            .custom-table td {{ font-size: 28px; /* 手机端八字大小 */ }}
        }}
        </style>

        <table class="custom-table">
            <thead>
                <tr>
                    <th>年柱</th>
                    <th>月柱</th>
                    <th>日柱</th>
                    <th>时柱</th>
                </tr>
            </thead>
            <tbody>
                <tr>
                    <td>{r['八字'][0]}</td>
                    <td>{r['八字'][1]}</td>
                    <td>{r['八字'][2]}</td>
                    <td>{r['八字'][3]}</td>
                </tr>
            </tbody>
        </table>
        """
        st.markdown(pillars_html, unsafe_allow_html=True)

        # 以下内容保持不变
        st.markdown(f"**公历**：{r['公历']}")
        st.markdown(f"**农历**：{r['农历']}")
        st.markdown(f"**生肖**：{r['生肖']}　**时辰**：{r['时辰']}　**日干**：{r['日干']}")
        st.markdown(
            f"**五行**：金{r['五行']['金']} 木{r['五行']['木']} 水{r['五行']['水']} 火{r['五行']['火']} 土{r['五行']['土']}")

# ===================== 完整Tabs模块（修复底部高亮同步） =====================
st.markdown("---")
tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab9 = st.tabs(["📆 万年历", "💰 八字论财", "🌀 八字合盘", "🔍 多盘对比", "☯️ 排盘", "⏱️ 吉日", "🏮 风水", "📄 解读", "📜 说明书"])
with tab1:
    d = st.date_input("选择日期", datetime.now(), min_value=datetime(1900, 1, 1), max_value=datetime(2100, 12, 31))

    # 【新增】查询完整黄历数据（五行、廿八宿、星期、十二宫辰、红砂）
    def get_full_calendar_data(solar_date_str: str):
        try:
            conn = sqlite3.connect(resource_path("bazi_calendar.db"), timeout=10)
            cursor = conn.cursor()
            cursor.execute('''
                SELECT 五行, 星期, 廿八宿, 十二宫辰, 红砂 
                FROM calendar WHERE 国历 = ? LIMIT 1
            ''', (solar_date_str,))
            row = cursor.fetchone()
            conn.close()
            if row:
                return {
                    "五行": row[0] or "",
                    "星期": row[1] or "",
                    "廿八宿": row[2] or "",
                    "十二宫辰": row[3] or "",
                    "红砂": row[4] if row[4] and row[4].strip() else "否"
                }
            else:
                return {"五行":"","星期":"","廿八宿":"","十二宫辰":"","红砂":"否"}
        except Exception:
            return {"五行":"","星期":"","廿八宿":"","十二宫辰":"","红砂":"否"}

    if st.button("查询万年历"):
        s = d.strftime("%Y-%m-%d")
        lu = solar_to_lunar_from_db(s)
        n, y, r = query_db_ganzhi(s)
        cal = get_full_calendar_data(s)

        st.write(f"📅 公历：{s}")
        st.write(f"📅 农历：{lu['农历完整信息']}")
        st.write(f"📅 星期：{cal['星期']}　🐉 生肖：{lu['生肖']}")
        st.write(f"🪶 年柱:{n}　月柱:{y}　日柱:{r}")
        st.write(f"🔥 五行：{cal['五行']}　🩸 红砂：{cal['红砂']}")
        st.write(f"🏮 廿八宿：{cal['廿八宿']}　🏯 十二宫辰：{cal['十二宫辰']}")
with tab2:
    if "bazi_result" in st.session_state and st.session_state.bazi_result:
        r = st.session_state.bazi_result
        ri_gan = r["日干"]
        w = r["五行"]

        # ===================== 显示姓名 + 性别 =====================
        # 读取顶部输入的姓名、性别
        show_name = name if name.strip() != "" else "命主"
        show_title = f"{show_name}{gender}"
        st.markdown(f"####    {show_title} · 八字财运")

        # ===================== 【传统正宗】按日主取财 =====================
        cai_wuxing = ""
        cai_name = ""
        if ri_gan in ("甲", "乙"):
            cai_wuxing = "土"
            cai_name = "土财（木克土为财）"
        elif ri_gan in ("丙", "丁"):
            cai_wuxing = "金"
            cai_name = "金财（火克金为财）"
        elif ri_gan in ("戊", "己"):
            cai_wuxing = "水"
            cai_name = "水财（土克水为财）"
        elif ri_gan in ("庚", "辛"):
            cai_wuxing = "木"
            cai_name = "木财（金克木为财）"
        elif ri_gan in ("壬", "癸"):
            cai_wuxing = "火"
            cai_name = "火财（水克火为财）"

        cai_num = w[cai_wuxing]

        # 财星旺衰
        if cai_num >= 3:
            cai_status = "财星极旺"
        elif cai_num == 2:
            cai_status = "财星偏旺"
        elif cai_num == 1:
            cai_status = "财星偏弱"
        else:
            cai_status = "财星微弱"

        st.markdown(f"**日主**：{ri_gan}")
        st.markdown(f"**你的财星**：{cai_name}")
        st.markdown(f"**财星数量**：{cai_num} 个")
        st.markdown(f"**财运状态**：{cai_status}")
        # st.markdown("---")
        st.markdown("""
**正财**：固定收入、薪资、稳定所得  
**偏财**：投资、外快、生意、意外之财  
""")
    else:
        st.warning("请先完成排盘，再查看财运分析")
with tab3:
    # st.markdown("#### 双人八字合盘")
    col_a, col_b = st.columns(2)
    with col_a:
        a_date = st.date_input("A公历生日", key="a_date")
        a_shichen = st.selectbox("A出生时辰", SHICHEN_DETAIL, key="a_shi")
        a_shichen_name = a_shichen.split(" ")[0]
    with col_b:
        b_date = st.date_input("B公历生日", key="b_date")
        b_shichen = st.selectbox("B出生时辰", SHICHEN_DETAIL, key="b_shi")
        b_shichen_name = b_shichen.split(" ")[0]

    # --------------------------
    # 合盘核心工具函数
    # --------------------------
    def get_zhi_relation(zhi1, zhi2):
        liuhe = ["子丑","寅亥","卯戌","辰酉","巳申","午未"]
        liuchong = ["子午","丑未","寅申","卯酉","辰戌","巳亥"]
        liuhai = ["子未","丑午","寅巳","卯辰","申亥","酉戌"]
        s1 = zhi1+zhi2
        s2 = zhi2+zhi1
        if s1 in liuhe or s2 in liuhe: return "六合"
        if s1 in liuchong or s2 in liuchong: return "六冲"
        if s1 in liuhai or s2 in liuhai: return "六害"
        return "无"

    def tian_gan_he(g1, g2):
        he_list = [("甲","己"),("乙","庚"),("丙","辛"),("丁","壬"),("戊","癸")]
        return (g1,g2) in he_list or (g2,g1) in he_list

    def ri_gan_wuxing(gan):
        if gan in "甲乙": return "木"
        if gan in "丙丁": return "火"
        if gan in "戊己": return "土"
        if gan in "庚辛": return "金"
        if gan in "壬癸": return "水"
        return ""

    def wuxing_shengke(w1, w2):
        sheng = {"木":"火","火":"土","土":"金","金":"水","水":"木"}
        ke = {"木":"土","土":"水","水":"火","火":"金","金":"木"}
        if sheng[w1] == w2: return "相生"
        if ke[w1] == w2: return "相克"
        return "相同"

    def get_spouse_star(ri_gan, gender):
        if gender in ("先生","男"):
            return "土" if ri_gan in "甲乙" else "金" if ri_gan in "丙丁" else "水" if ri_gan in "戊己" else "木" if ri_gan in "庚辛" else "火"
        else:
            return "火" if ri_gan in "甲乙" else "土" if ri_gan in "丙丁" else "金" if ri_gan in "戊己" else "水" if ri_gan in "庚辛" else "木"

    if st.button("开始婚姻合盘"):
        a_data = BaziCalculator.generate_bazi(a_date.strftime("%Y-%m-%d"), a_shichen_name)
        b_data = BaziCalculator.generate_bazi(b_date.strftime("%Y-%m-%d"), b_shichen_name)
        a_bazi = a_data["八字"]
        b_bazi = b_data["八字"]
        a_nian, a_yue, a_ri, a_shi = a_bazi
        b_nian, b_yue, b_ri, b_shi = b_bazi
        a_rg = a_data["日干"]
        b_rg = b_data["日干"]
        a_wx = a_data["五行"]
        b_wx = b_data["五行"]

        score = 0
        items = []

        # 1）生肖（年支）关系
        zhi_rel = get_zhi_relation(a_nian[1], b_nian[1])
        if zhi_rel == "六合":
            score += 15
            items.append("生肖六合｜缘分深 +15")
        elif zhi_rel == "六冲":
            score -= 10
            items.append("生肖六冲｜易矛盾 -10")
        elif zhi_rel == "六害":
            score -= 5
            items.append("生肖六害｜暗耗 -5")
        else:
            score += 5
            items.append("生肖平和 +5")

        # 2）夫妻宫（日支）
        fg_rel = get_zhi_relation(a_ri[1], b_ri[1])
        if fg_rel == "六合":
            score += 20
            items.append("夫妻宫六合｜极佳 +20")
        elif fg_rel == "六冲":
            score -= 20
            items.append("夫妻宫六冲｜不稳 -20")
        elif fg_rel == "六害":
            score -= 10
            items.append("夫妻宫六害｜内耗 -10")
        else:
            score += 8
            items.append("夫妻宫平和 +8")

        # 3）日主天干合
        if tian_gan_he(a_rg, b_rg):
            score += 18
            items.append("日主天干相合｜情投意合 +18")
        else:
            w1 = ri_gan_wuxing(a_rg)
            w2 = ri_gan_wuxing(b_rg)
            rel = wuxing_shengke(w1, w2)
            if rel == "相生":
                score += 12
                items.append(f"日主相生｜滋养 +12")
            elif rel == "相克":
                score -= 8
                items.append(f"日主相克｜摩擦 -8")
            else:
                score += 4
                items.append("日主五行相同 +4")

        # 4）五行互补
        hubu = 0
        for x in ["金","木","水","火","土"]:
            if (a_wx[x]>0) != (b_wx[x]>0):
                hubu += 1
        score += hubu * 6
        items.append(f"五行互补 {hubu}/5 项 +{hubu*6}")

        # 5）夫妻星匹配
        a_star = get_spouse_star(a_rg, "先生")
        b_star = get_spouse_star(b_rg, "女士")
        a_has = a_wx[a_star] > 0
        b_has = b_wx[b_star] > 0
        if a_has and b_has:
            score += 15
            items.append("夫妻星皆有｜婚配佳 +15")
        else:
            items.append("夫妻星偏弱")

        # 总分区间
        score = max(score, 0)
        score = min(score, 100)
        if score >= 80:
            level = "上等婚配｜天生一对"
            color = "#28a745"
        elif score >= 60:
            level = "上等婚配｜和谐美满"
            color = "#17a2b8"
        elif score >= 40:
            level = "中等婚配｜可长久"
            color = "#ffc107"
        else:
            level = "普通婚配｜多包容"
            color = "#dc3545"

        # 展示结果
        st.markdown("---")
        st.markdown(f"**A方八字**：{a_data['八字_str']}")
        st.markdown(f"**B方八字**：{b_data['八字_str']}")
        st.markdown(f"### 合盘总分：<font color='{color}'>{score}分</font>｜{level}", unsafe_allow_html=True)
        st.markdown("#### 评分明细")
        for txt in items:
            st.markdown(f"- {txt}")
with tab4:
    # st.markdown("#### 多盘对比")
    if "duopan_list" not in st.session_state:
        st.session_state.duopan_list = []

    col_btn_add, col_btn_clear, col_btn_compare = st.columns(3)
    with col_btn_add:
        if st.button("添加当前八字", use_container_width=True):
            if "bazi_result" in st.session_state and st.session_state.bazi_result:
                st.session_state.duopan_list.append(st.session_state.bazi_result)
                st.success("已添加到对比列表")
    with col_btn_clear:
        if st.button("清空当前列表", use_container_width=True):
            st.session_state.duopan_list = []
            st.success("已清空")
    with col_btn_compare:
        if st.button("开始多盘对比", use_container_width=True):
            if len(st.session_state.duopan_list) < 2:
                st.warning("至少添加2个八字才能对比")
            else:
                # =============== 专业多盘分析核心逻辑 ===============
                all_wuxing = []
                all_rigan = []
                total_count = len(st.session_state.duopan_list)

                for d in st.session_state.duopan_list:
                    all_wuxing.append(d["五行"])
                    all_rigan.append(d["日干"])

                # 1. 计算五行总和、平均值
                sum_wuxing = {"金":0,"木":0,"水":0,"火":0,"土":0}
                avg_wuxing = {"金":0.0,"木":0.0,"水":0.0,"火":0.0,"土":0.0}
                min_wuxing = {"金":99,"木":99,"水":99,"火":99,"土":99}
                max_wuxing = {"金":0,"木":0,"水":0,"火":0,"土":0}

                for wd in all_wuxing:
                    for k in sum_wuxing:
                        sum_wuxing[k] += wd[k]
                        if wd[k] < min_wuxing[k]:
                            min_wuxing[k] = wd[k]
                        if wd[k] > max_wuxing[k]:
                            max_wuxing[k] = wd[k]

                for k in sum_wuxing:
                    avg_wuxing[k] = sum_wuxing[k] / total_count

                # 2. 找最旺、最弱五行
                sorted_wuxing = sorted(avg_wuxing.items(), key=lambda x:x[1], reverse=True)
                top1 = sorted_wuxing[0][0]
                top2 = sorted_wuxing[1][0]
                low1 = sorted_wuxing[-1][0]
                low2 = sorted_wuxing[-2][0]

                # 3. 群体整体判断
                total_score = sum(sum_wuxing.values())
                avg_per = total_score / total_count
                if avg_per >= 7:
                    group_type = "整体偏旺｜行动力强、易冲动"
                elif avg_per <= 4:
                    group_type = "整体偏弱｜温和内敛、易劳累"
                else:
                    group_type = "五行均衡｜稳定协调、适配性强"

                # 4. 适合行业
                industry_map = {
                    "木":"教育、文化、创意、医疗",
                    "火":"能源、餐饮、互联网、娱乐",
                    "土":"地产、建筑、农业、保险",
                    "金":"金融、法律、金属、管理",
                    "水":"运输、贸易、传媒、旅游"
                }
                best_industry = industry_map.get(top1, "")
                need_industry = industry_map.get(low1, "")

                # =============== 展示结果 ===============
                st.markdown("---")
                st.success(f"✅ 群体分析完成｜共 {total_count} 个命盘")
                st.markdown(f"**🏛️ 群体整体类型**：{group_type}")
                st.markdown(f"**🔥 最旺五行**：{top1}、{top2}")
                st.markdown(f"**💧 最弱/最缺五行**：{low1}、{low2}")
                st.markdown(f"**📌 群体适合行业**：{best_industry}")
                st.markdown(f"**📌 建议补充五行**：{need_industry}")
                st.markdown("---")
                st.markdown("#### 五行详细数据")
                for k in ["金","木","水","火","土"]:
                    st.markdown(f"**{k}**：均值 {avg_wuxing[k]:.1f}　范围 {min_wuxing[k]} ~ {max_wuxing[k]}")

    # 显示已添加列表
    if st.session_state.duopan_list:
        st.markdown("---")
        st.markdown("**已添加对比命盘**")
        for i, data in enumerate(st.session_state.duopan_list):
            st.markdown(f"{i+1}.{data['八字_str']}{data['农历']}")
    else:
        st.info("先排盘 → 添加到对比列表 → 开始分析")
with tab5:
    if st.button("排盘首页",use_container_width=True):
        st.session_state.bottom_nav_active = "排盘"
        # 这里执行排盘功能代码
        st.rerun()
with tab6:
    if st.button("查询吉日",use_container_width=True):
        st.session_state.bottom_nav_active = "吉日"
        # 这里执行吉日功能代码
        st.rerun()
with tab7:
    if st.button("查询风水",use_container_width=True):
        st.session_state.bottom_nav_active = "风水"
        # 这里执行风水功能代码
        st.rerun()
with tab8:
    if st.button("deepseek 八 字 解 读",use_container_width=True):
        st.session_state.bottom_nav_active = "解读"
        # 这里执行解读功能代码
        st.rerun()
with tab9:
    st.markdown("<div style='text-align:center; margin-top:10px;'><h3>📜 真命盘说明书</h3></div>", unsafe_allow_html=True)
    st.markdown("---")
    st.markdown("""
### 一、开发作者简介
吴勇，计算机网络高级工程师、网络信息安全高级技师、省级名师工作室成员、省级优秀指导教师、高级双师型教师。爱好计算机编程与传统文化数字化领域，擅长将正统民俗命理逻辑与软件工程规范结合，打造严谨、稳定、易用的命理工具系统。

### 二、真命盘系统介绍
真命盘是一套基于正统子平八字 + 传统黄历的命理数字化系统。系统以权威万年历数据库为核心，不做任何自定义命理推算，所有干支、神煞、吉凶均取自预制黄历数据，严格遵循传统命理规则，实现：八字排盘、专业择日、八字合盘、多盘对比、民俗风水、AI 深度解读、报告导出七大核心功能。系统特点：纯数据库读取、无运算误差、逻辑严谨、界面简洁、PC / 移动端无缝互通，适合命理学习、民俗参考、日常择吉使用。

### 三、择日功能介绍
真命盘择日功能严格遵循传统择吉文化 “先避凶、后趋吉” 的核心思维，以正统黄历数据为依据，层层筛查，确保所选日子符合民俗规范。

#### 1. 核心筛查逻辑
第一步：避大凶（红砂日）
传统民俗 “红砂岀，百事忌”，系统直接读取数据库红砂字段，自动排除所有红砂凶日。

第二步：避冲煞
读取日干支，排除冲命主生肖、冲命主日柱的日子，确保不与自身命格相冲。

第三步：避杂煞
排除四离四绝、三娘煞、月破日等传统大忌日子，气场紊乱之日不做大事。

第四步：选黄道
仅保留青龙、明堂、金匮、天德、玉堂、司命六大黄道吉日。

第五步：事项匹配
按事项匹配日支：
开业 / 出行：寅申巳亥日
嫁娶 / 订婚：子午卯酉日
入宅 / 动土：辰戌丑未日
安葬：子丑辰未申酉日

第六步：多人适配
嫁娶 / 订婚限 2 人，入宅 / 动工支持多人，不冲任何参与者，支持双胞胎相同排盘。

#### 2. 分级结果
太阳吉日：最高等级，带太阳星吉时
完美吉日：带吉神加持
平安吉日：基础可用吉日

#### 3. 产品优势
纯数据库：无推算、无误差，结果稳定
民俗正统：完全贴合传统择日思维
场景完善：覆盖 12 类常用事项。

**特别提醒**：命理分析仅为传统文化趣味解读，人生走向取决于个人努力与选择，请理性看待，勿盲从迷信！
""")
# ===================== APP.py 终极修复：支持双胞胎+嫁娶限2人+对齐main + 新增验证吉日 =====================
if st.session_state.bottom_nav_active == "吉日":
    st.markdown("<div style='text-align:center;margin-top:20px;'><h3>📅 专业择日系统</h3></div>", unsafe_allow_html=True)
    if "bazi_result" not in st.session_state or not st.session_state.bazi_result:
        st.warning("⚠️ 请先排盘再查询吉日")
        st.stop()

    r = st.session_state.bazi_result
    # ===================== 新增：模式切换（查询吉日 / 验证公历吉日） =====================
    check_mode = st.radio("", ["查询5年内吉日", "验证*公历吉日"], horizontal=True, label_visibility="collapsed")

    jiri_type = st.radio(
        "", ["开业择日", "出行择日", "上任择日", "祈福择日", "修灶择日", "财门择日",
             "嫁娶择日", "订婚择日", "入宅择日", "动工择日", "动土择日", "安葬择日"],
        horizontal=True, label_visibility="collapsed"
    )
    # 严格分类
    single_types = ["开业择日", "出行择日", "上任择日", "祈福择日", "修灶择日", "财门择日"]
    double_types = ["嫁娶择日", "订婚择日"]  # 限2人
    multi_types = ["入宅择日", "动工择日", "动土择日", "安葬择日"]
    is_single = jiri_type in single_types
    is_double = jiri_type in double_types
    is_multi = jiri_type in multi_types

    # 双人/多人界面（原有逻辑不变）
    if not is_single:
        st.markdown("---")
        if is_double:
            st.markdown("🔴 **嫁娶/订婚择日：仅限 2 人（男女双方），支持同盘**")
        else:
            st.markdown("🔹 **多人择日：可添加多人，支持重复添加同盘**")
        col_add, col_clear, col_start = st.columns(3)
        with col_add:
            if st.button("➕ 添加当前八字", use_container_width=True):
                if "jiri_list" not in st.session_state:
                    st.session_state.jiri_list = []
                st.session_state.jiri_list.append(r)
                if is_double:
                    if len(st.session_state.jiri_list) > 2:
                        st.session_state.jiri_list = st.session_state.jiri_list[-2:]
                        st.warning("⚠️ 嫁娶/订婚仅限2人，已自动保留最后2个")
                    else:
                        st.success("✅ 已添加")
                else:
                    st.success("✅ 已添加")
        with col_clear:
            if st.button("🧹 清空列表", use_container_width=True):
                st.session_state.jiri_list = []
                st.success("🗑️ 已清空")
        # 显示列表
        if st.session_state.get("jiri_list"):
            st.markdown("**📋 参与择日八字**")
            for i, item in enumerate(st.session_state.jiri_list):
                st.write(f"{i + 1}. {item['八字_str']}")

    # ===================== 核心：验证指定公历吉日（仅公历） =====================
    if check_mode == "验证*公历吉日":
        st.markdown("---")
        st.markdown("🔍 验证指定公历日期是否为吉日")
        # 日期选择器（仅公历）
        check_date = st.date_input("请输入要验证的公历日期", datetime.now(),
                                   min_value=datetime(1900, 1, 1),
                                   max_value=datetime(2100, 12, 31),
                                   label_visibility="collapsed")
        check_date_str = check_date.strftime("%Y-%m-%d")

        # 验证按钮
        if st.button("✅ 开始验证吉日", use_container_width=True):
            with st.spinner("正在验证中..."):
                # 原有常量（完全复用）
                chong_map = {"子": "午", "丑": "未", "寅": "申", "卯": "酉", "辰": "戌", "巳": "亥", "午": "子",
                             "未": "丑", "申": "寅", "酉": "卯", "戌": "辰", "亥": "巳"}
                sx_map = {"鼠": "子", "牛": "丑", "虎": "寅", "兔": "卯", "龙": "辰", "蛇": "巳", "马": "午",
                          "羊": "未", "猴": "申", "鸡": "酉", "狗": "戌", "猪": "亥"}
                sili_jue = [(3, 20), (6, 21), (9, 23), (12, 22), (2, 4), (5, 6), (8, 8), (11, 8)]
                sanniang = [3, 7, 13, 18, 22, 27]
                sun_star_table = {
                    "甲子": ([6, 15, 24], "未"), "乙丑": ([1, 10, 19, 28], "申"), "丙寅": ([1, 10, 19, 28], "辰"),
                    "丁卯": ([8, 17, 28], "申"),
                    "戊辰": ([1, 10, 19, 28], "卯"), "己巳": ([1, 10, 19, 28], "未"), "庚午": ([5, 14, 23], "申"),
                    "辛未": ([6, 15, 24], "辰"),
                    "壬申": ([5, 14, 23], "申"), "癸酉": ([9, 18, 27], "卯"), "甲戌": ([7, 16, 25], "未"),
                    "乙亥": ([4, 13, 22], "申"),
                    "丙子": ([9, 18, 27], "辰"), "丁丑": ([3, 12, 21, 30], "申"), "戊寅": ([3, 12, 21, 30], "卯"),
                    "己卯": ([8, 17, 28], "未"),
                    "庚辰": ([1, 10, 19, 28], "申"), "辛巳": ([1, 10, 19, 28], "辰"), "壬午": ([6, 15, 24], "申"),
                    "癸未": ([5, 14, 23], "卯"),
                    "甲申": ([5, 14, 23], "未"), "乙酉": ([2, 11, 20, 29], "申"), "丙戌": ([4, 13, 22], "辰"),
                    "丁亥": ([7, 16, 25], "申"),
                    "戊子": ([2, 11, 20, 29], "卯"), "己丑": ([3, 12, 21, 30], "未"), "庚寅": ([8, 17, 28], "申"),
                    "辛卯": ([3, 12, 21, 30], "辰"),
                    "壬辰": ([3, 12, 21, 30], "申"), "癸巳": ([3, 12, 21, 30], "卯"), "甲午": ([2, 11, 20, 29], "未"),
                    "乙未": ([4, 13, 22], "申"),
                    "丙申": ([4, 13, 22], "辰"), "丁酉": ([9, 18, 27], "申"), "戊戌": ([5, 14, 23], "卯"),
                    "己亥": ([5, 14, 23], "未"),
                    "庚子": ([5, 14, 23], "申"), "辛丑": ([1, 10, 19, 28], "辰"), "壬寅": ([1, 10, 19, 28], "申"),
                    "癸卯": ([8, 17, 28], "卯"),
                    "甲辰": ([1, 10, 19, 28], "未"), "乙巳": ([1, 10, 19, 28], "申"), "丙午": ([5, 14, 23], "辰"),
                    "丁未": ([5, 14, 23], "申"),
                    "戊申": ([5, 14, 23], "卯"), "己酉": ([9, 17, 28], "未"), "庚戌": ([4, 13, 22], "申"),
                    "辛亥": ([4, 13, 22], "辰"),
                    "壬子": ([2, 11, 20, 29], "申"), "癸丑": ([3, 12, 21, 30], "卯"), "甲寅": ([3, 12, 21, 30], "未"),
                    "乙卯": ([3, 12, 21, 30], "申"),
                    "丙辰": ([8, 17, 28], "辰"), "丁巳": ([3, 12, 21, 30], "申"), "戊午": ([2, 11, 20, 29], "卯"),
                    "己未": ([7, 16, 25], "未"),
                    "庚申": ([4, 13, 22], "申"), "辛酉": ([2, 11, 20, 29], "辰"), "壬戌": ([5, 14, 23], "申"),
                    "癸亥": ([5, 14, 23], "卯"),
                }


                # 原有函数（完全复用）
                def get_huangdao(zhi):
                    hd = {"子": "青龙", "丑": "明堂", "寅": "天刑", "卯": "朱雀", "辰": "金匮", "巳": "天德",
                          "午": "白虎", "未": "玉堂", "申": "天牢", "酉": "玄武", "戌": "司命", "亥": "勾陈"}
                    return hd.get(zhi) in ["青龙", "明堂", "金匮", "天德", "玉堂", "司命"]


                def is_yue_po(month, zhi):
                    yz = ["", "寅", "卯", "辰", "巳", "午", "未", "申", "酉", "戌", "亥", "子", "丑"]
                    return chong_map.get(yz[month], "") == zhi


                def match_type(zhi, t):
                    if t in ["开业择日", "上任择日", "出行择日", "财门择日"]: return zhi in ["寅", "申", "巳", "亥"]
                    if t in ["嫁娶择日", "订婚择日", "修灶择日"]: return zhi in ["子", "午", "卯", "酉"]
                    if t in ["入宅择日", "祈福择日", "动工择日", "动土择日"]: return zhi in ["辰", "戌", "丑", "未"]
                    if t in ["安葬择日"]: return zhi in ["子", "丑", "辰", "未", "申", "酉"]
                    return True


                def check_god(dg, dz, yue_zhi_main, nian_gan):
                    lu = {"甲": "寅", "乙": "卯", "丙": "巳", "丁": "午", "戊": "巳", "己": "午", "庚": "申",
                          "辛": "酉", "壬": "亥", "癸": "子"}
                    tx = {"寅": "戌", "卯": "亥", "辰": "子", "巳": "丑", "午": "寅", "未": "卯", "申": "辰",
                          "酉": "巳", "戌": "午", "亥": "未", "子": "申", "丑": "酉"}
                    ye = {"寅": "亥", "卯": "子", "辰": "丑", "巳": "寅", "午": "卯", "未": "辰", "申": "巳",
                          "酉": "午", "戌": "未", "亥": "申", "子": "酉", "丑": "戌"}
                    td = {"甲": ["寅", "午"], "乙": ["申", "子"], "丙": ["卯", "亥"], "丁": ["巳", "丑"],
                          "戊": ["卯", "亥"], "己": ["巳", "丑"], "庚": ["子", "申"], "辛": ["寅", "午"],
                          "壬": ["巳", "丑"], "癸": ["卯", "亥"]}
                    return dz == lu.get(dg, "") or dz == tx.get(yue_zhi_main, "") or dz == ye.get(yue_zhi_main,
                                                                                                  "") or dz in td.get(
                        nian_gan, [])


                # 原有参数
                ri_zhi = r["八字"][2][1]
                shengxiao = r["生肖"]
                yue_zhi_main = r["八字"][1][1]
                nian_gan = r["八字"][0][0]
                m, d = check_date.month, check_date.day
                reason = []
                is_valid = True

                # 数据库查询（仅验证指定日期）
                try:
                    conn = sqlite3.connect(resource_path("bazi_calendar.db"), timeout=10)
                    cursor = conn.cursor()
                    cursor.execute("SELECT 纳音,农历日,红砂 FROM calendar WHERE 国历 LIKE ? LIMIT 1",
                                   (check_date_str + "%",))
                    res = cursor.fetchone()
                    conn.close()
                    if not res:
                        st.error("❌ 未查询到该日期数据")
                        st.stop()
                    day_gz, lunar_day, hongsha = res
                    if len(day_gz) != 2:
                        st.error("❌ 日期干支数据异常")
                        st.stop()
                    dg, dz = day_gz[0], day_gz[1]

                    # 全套择吉验证（和原有规则完全一致）
                    if hongsha and hongsha.strip() == "红砂":
                        reason.append("❌ 红砂日（大凶）")
                        is_valid = False
                    # 冲煞校验
                    sx_zhi = sx_map.get(shengxiao, "")
                    if dz == chong_map.get(sx_zhi, "") or dz == chong_map.get(ri_zhi, ""):
                        reason.append("❌ 冲命主生肖/日柱")
                        is_valid = False
                    # 多人冲煞校验
                    if not is_single:
                        for b in st.session_state.get("jiri_list", []):
                            try:
                                if dz == chong_map.get(b["八字"][2][1], ""):
                                    reason.append("❌ 冲参与人八字")
                                    is_valid = False
                                    break
                            except:
                                continue
                    if (m, d) in sili_jue:
                        reason.append("❌ 四离四绝日")
                        is_valid = False
                    if d in sanniang:
                        reason.append("❌ 三娘煞日")
                        is_valid = False
                    if is_yue_po(m, dz):
                        reason.append("❌ 月破日")
                        is_valid = False
                    if not get_huangdao(dz):
                        reason.append("❌ 非黄道吉日")
                        is_valid = False
                    if not match_type(dz, jiri_type):
                        reason.append("❌ 不匹配该事项")
                        is_valid = False

                    # 太阳星/吉神判断
                    first_str = check_date.replace(day=1).strftime("%Y-%m-%d")
                    conn = sqlite3.connect(resource_path("bazi_calendar.db"), timeout=10)
                    cursor = conn.cursor()
                    cursor.execute("SELECT 纳音 FROM calendar WHERE 国历 LIKE ? LIMIT 1", (first_str + "%",))
                    frow = cursor.fetchone()
                    conn.close()
                    first_gz = frow[0] if frow else ""
                    is_sun = False
                    sun_time = ""
                    if first_gz in sun_star_table and lunar_day in sun_star_table[first_gz][0]:
                        is_sun = True
                        sun_time = sun_star_table[first_gz][1]
                    has_god = check_god(dg, dz, yue_zhi_main, nian_gan)

                    # 输出验证结果
                    st.markdown("---")
                    st.markdown(f"📅 验证日期：{check_date_str}({day_gz})")
                    st.markdown(f"📌 验证事项：{jiri_type}")
                    if is_valid:
                        st.success("✅ 该日期为**合格吉日**")
                        if is_sun:
                            st.markdown("<center><font color=#D4AF37>☀️ 太阳吉日</font></center>",
                                        unsafe_allow_html=True)
                            st.write(f"太阳吉时：{sun_time}")
                        elif has_god:
                            st.markdown("<center><font color=#ff6666>🌟 完美吉日</font></center>",
                                        unsafe_allow_html=True)
                        else:
                            st.markdown("<center><font color=#007fff>🛡️ 平安吉日</font></center>",
                                        unsafe_allow_html=True)
                    else:
                        st.error("❌ 该日期**不是吉日**")
                        st.markdown("#### 不合格原因：")
                        for txt in reason:
                            st.write(txt)
                except Exception as e:
                    st.error(f"❌ 验证失败：{str(e)}")
        st.stop()

    # ===================== 原有：查询5年内吉日逻辑（完全不变） =====================
    allow_query = False
    if is_single:
        allow_query = True
    elif is_double:
        allow_query = len(st.session_state.get("jiri_list", [])) == 2
    else:
        allow_query = len(st.session_state.get("jiri_list", [])) >= 2

    if st.button("🔍 查询5年内吉日", use_container_width=True):
        if not allow_query:
            if is_double:
                st.warning("⚠️ 嫁娶/订婚必须添加 **恰好2个八字**！")
            else:
                st.warning("⚠️ 请先添加至少 **2个八字**！")
            st.stop()
        with st.spinner("正在筛选吉日..."):
            chong_map = {"子": "午", "丑": "未", "寅": "申", "卯": "酉", "辰": "戌", "巳": "亥", "午": "子", "未": "丑",
                         "申": "寅", "酉": "卯", "戌": "辰", "亥": "巳"}
            sx_map = {"鼠": "子", "牛": "丑", "虎": "寅", "兔": "卯", "龙": "辰", "蛇": "巳", "马": "午", "羊": "未",
                      "猴": "申", "鸡": "酉", "狗": "戌", "猪": "亥"}
            sili_jue = [(3, 20), (6, 21), (9, 23), (12, 22), (2, 4), (5, 6), (8, 8), (11, 8)]
            sanniang = [3, 7, 13, 18, 22, 27]
            sun_star_table = {
                "甲子": ([6, 15, 24], "未"), "乙丑": ([1, 10, 19, 28], "申"), "丙寅": ([1, 10, 19, 28], "辰"),
                "丁卯": ([8, 17, 28], "申"),
                "戊辰": ([1, 10, 19, 28], "卯"), "己巳": ([1, 10, 19, 28], "未"), "庚午": ([5, 14, 23], "申"),
                "辛未": ([6, 15, 24], "辰"),
                "壬申": ([5, 14, 23], "申"), "癸酉": ([9, 18, 27], "卯"), "甲戌": ([7, 16, 25], "未"),
                "乙亥": ([4, 13, 22], "申"),
                "丙子": ([9, 18, 27], "辰"), "丁丑": ([3, 12, 21, 30], "申"), "戊寅": ([3, 12, 21, 30], "卯"),
                "己卯": ([8, 17, 28], "未"),
                "庚辰": ([1, 10, 19, 28], "申"), "辛巳": ([1, 10, 19, 28], "辰"), "壬午": ([6, 15, 24], "申"),
                "癸未": ([5, 14, 23], "卯"),
                "甲申": ([5, 14, 23], "未"), "乙酉": ([2, 11, 20, 29], "申"), "丙戌": ([4, 13, 22], "辰"),
                "丁亥": ([7, 16, 25], "申"),
                "戊子": ([2, 11, 20, 29], "卯"), "己丑": ([3, 12, 21, 30], "未"), "庚寅": ([8, 17, 28], "申"),
                "辛卯": ([3, 12, 21, 30], "辰"),
                "壬辰": ([3, 12, 21, 30], "申"), "癸巳": ([3, 12, 21, 30], "卯"), "甲午": ([2, 11, 20, 29], "未"),
                "乙未": ([4, 13, 22], "申"),
                "丙申": ([4, 13, 22], "辰"), "丁酉": ([9, 18, 27], "申"), "戊戌": ([5, 14, 23], "卯"),
                "己亥": ([5, 14, 23], "未"),
                "庚子": ([5, 14, 23], "申"), "辛丑": ([1, 10, 19, 28], "辰"), "壬寅": ([1, 10, 19, 28], "申"),
                "癸卯": ([8, 17, 28], "卯"),
                "甲辰": ([1, 10, 19, 28], "未"), "乙巳": ([1, 10, 19, 28], "申"), "丙午": ([5, 14, 23], "辰"),
                "丁未": ([5, 14, 23], "申"),
                "戊申": ([5, 14, 23], "卯"), "己酉": ([9, 17, 28], "未"), "庚戌": ([4, 13, 22], "申"),
                "辛亥": ([4, 13, 22], "辰"),
                "壬子": ([2, 11, 20, 29], "申"), "癸丑": ([3, 12, 21, 30], "卯"), "甲寅": ([3, 12, 21, 30], "未"),
                "乙卯": ([3, 12, 21, 30], "申"),
                "丙辰": ([8, 17, 28], "辰"), "丁巳": ([3, 12, 21, 30], "申"), "戊午": ([2, 11, 20, 29], "卯"),
                "己未": ([7, 16, 25], "未"),
                "庚申": ([4, 13, 22], "申"), "辛酉": ([2, 11, 20, 29], "辰"), "壬戌": ([5, 14, 23], "申"),
                "癸亥": ([5, 14, 23], "卯"),
            }


            def get_huangdao(zhi):
                hd = {"子": "青龙", "丑": "明堂", "寅": "天刑", "卯": "朱雀", "辰": "金匮", "巳": "天德", "午": "白虎",
                      "未": "玉堂", "申": "天牢", "酉": "玄武", "戌": "司命", "亥": "勾陈"}
                return hd.get(zhi) in ["青龙", "明堂", "金匮", "天德", "玉堂", "司命"]


            def is_yue_po(month, zhi):
                yz = ["", "寅", "卯", "辰", "巳", "午", "未", "申", "酉", "戌", "亥", "子", "丑"]
                return chong_map.get(yz[month], "") == zhi


            def match_type(zhi, t):
                if t in ["开业择日", "上任择日", "出行择日", "财门择日"]: return zhi in ["寅", "申", "巳", "亥"]
                if t in ["嫁娶择日", "订婚择日", "修灶择日"]: return zhi in ["子", "午", "卯", "酉"]
                if t in ["入宅择日", "祈福择日", "动工择日", "动土择日"]: return zhi in ["辰", "戌", "丑", "未"]
                if t in ["安葬择日"]: return zhi in ["子", "丑", "辰", "未", "申", "酉"]
                return True


            def check_god(dg, dz, yue_zhi_main, nian_gan):
                lu = {"甲": "寅", "乙": "卯", "丙": "巳", "丁": "午", "戊": "巳", "己": "午", "庚": "申", "辛": "酉",
                      "壬": "亥", "癸": "子"}
                tx = {"寅": "戌", "卯": "亥", "辰": "子", "巳": "丑", "午": "寅", "未": "卯", "申": "辰", "酉": "巳",
                      "戌": "午", "亥": "未", "子": "申", "丑": "酉"}
                ye = {"寅": "亥", "卯": "子", "辰": "丑", "巳": "寅", "午": "卯", "未": "辰", "申": "巳", "酉": "午",
                      "戌": "未", "亥": "申", "子": "酉", "丑": "戌"}
                td = {"甲": ["寅", "午"], "乙": ["申", "子"], "丙": ["卯", "亥"], "丁": ["巳", "丑"],
                      "戊": ["卯", "亥"], "己": ["巳", "丑"], "庚": ["子", "申"], "辛": ["寅", "午"],
                      "壬": ["巳", "丑"], "癸": ["卯", "亥"]}
                return dz == lu.get(dg, "") or dz == tx.get(yue_zhi_main, "") or dz == ye.get(yue_zhi_main,
                                                                                              "") or dz in td.get(
                    nian_gan, [])


            ri_zhi = r["八字"][2][1]
            shengxiao = r["生肖"]
            yue_zhi_main = r["八字"][1][1]
            nian_gan = r["八字"][0][0]
            today = datetime.now()
            sun_best, perfect, safe = [], [], []
            conn = sqlite3.connect(resource_path("bazi_calendar.db"), timeout=10)
            cursor = conn.cursor()
            for i in range(1, 1826):
                dt = today + timedelta(days=i)
                date_str = dt.strftime("%Y-%m-%d")
                m, d = dt.month, dt.day
                cursor.execute("SELECT 纳音,农历日,红砂 FROM calendar WHERE 国历 LIKE ? LIMIT 1", (date_str + "%",))
                res = cursor.fetchone()
                if not res: continue
                day_gz, lunar_day, hongsha = res
                if len(day_gz) != 2: continue
                dg, dz = day_gz[0], day_gz[1]
                if hongsha and hongsha.strip() == "红砂": continue
                sx_zhi = sx_map.get(shengxiao, "")
                if dz == chong_map.get(sx_zhi, "") or dz == chong_map.get(ri_zhi, ""): continue
                if not is_single:
                    skip = False
                    for b in st.session_state.jiri_list:
                        try:
                            if dz == chong_map.get(b["八字"][2][1], ""):
                                skip = True
                                break
                        except:
                            continue
                    if skip: continue
                if (m, d) in sili_jue or d in sanniang or is_yue_po(m, dz): continue
                if not get_huangdao(dz) or not match_type(dz, jiri_type): continue
                first_str = dt.replace(day=1).strftime("%Y-%m-%d")
                cursor.execute("SELECT 纳音 FROM calendar WHERE 国历 LIKE ? LIMIT 1", (first_str + "%",))
                frow = cursor.fetchone()
                first_gz = frow[0] if frow else ""
                is_sun, sun_time = False, ""
                if first_gz in sun_star_table and lunar_day in sun_star_table[first_gz][0]:
                    is_sun = True
                    sun_time = sun_star_table[first_gz][1]
                has_god = check_god(dg, dz, yue_zhi_main, nian_gan)
                line = f"{date_str}({day_gz})"
                if has_god: line += "⭐吉"
                if is_sun:
                    sun_best.append(f"{line} 太阳吉时:{sun_time}")
                elif has_god:
                    perfect.append(line)
                else:
                    safe.append(line)
            conn.close()
            st.markdown("---")
            st.success(f"✅ {jiri_type} 筛选完成")
            if sun_best:
                st.markdown("<center><font color=#D4AF37>☀️ 太阳吉日</font></center>", unsafe_allow_html=True)
                for s in sun_best[:3]: st.write(s)
            if perfect:
                st.markdown("<center><font color=#ff6666>🌟 完美吉日</font></center>", unsafe_allow_html=True)
                for s in perfect[:5]: st.write(s)
            if safe:
                st.markdown("<center><font color=#007fff>🛡️ 平安吉日</font></center>", unsafe_allow_html=True)
                for s in safe[:10]: st.write(s)
            if not sun_best and not perfect and not safe:
                st.info("未找到符合条件的吉日")

# ===================== 独立风水页面（PC版原版完整移植·不影响任何功能） =====================
if st.session_state.bottom_nav_active == "风水":
    st.markdown("<div style='text-align:center; margin-top:20px;'><h3>🧭 民俗风水·专属</h3></div>",
                unsafe_allow_html=True)
    if "bazi_result" not in st.session_state or not st.session_state.bazi_result:
        st.warning("⚠️ 请先在排盘页完成排盘，再查看专属风水")
    else:
        r = st.session_state.bazi_result
        gender = st.session_state.get("gender", "先生")
        shengxiao = r["生肖"]
        ri_gan = r["日干"]
        bazi_str = r["八字_str"]


        # ============== 原版PC风水核心计算逻辑（100%复刻） ==============
        def get_fengshui_advice():
            try:
                # 1. 生肖五行
                wx = SHENGXIAO_WUXING.get(shengxiao, "土")
                # 2. 八宅命卦：男坎女坤（商用简化正统版）
                is_male = gender == "先生"
                minggua = "坎" if is_male else "坤"
                zhai_type = "东四宅" if minggua in ["坎", "离", "震", "巽"] else "西四宅"
                # 3. 吉方位
                caifeng = YONGSHEN_FANGWEI.get(wx, "北方")
                shiye = caifeng
                wenchang = "东南"
                taohua = "南方" if wx in ["木", "火"] else "西方"
                # 4. 吉利楼层
                he_lou = LOU_CENG_WUXING.get(wx, [1, 6])
                # 5. 幸运色
                se = {"木": "青绿", "火": "红紫", "土": "黄棕", "金": "白金", "水": "蓝黑"}[wx]
                # 6. 输出文本
                text = f"【一、命卦与宅卦】\n"
                text += f"卦象：{minggua}卦｜{'东四命' if zhai_type == '东四宅' else '西四命'}\n"
                text += f"适合住宅：{zhai_type}\n\n"
                text += f"【二、专属吉方位】\n"
                text += f"求财吉位：{caifeng}｜事业吉位：{shiye}\n"
                text += f"文昌吉位：{wenchang}｜桃花吉位：{taohua}\n\n"
                text += f"【三、八大方位吉凶】\n"
                j = BAZHAI_JIXIONG[:4]
                x = BAZHAI_JIXIONG[4:]
                text += "吉位：" + "｜".join(j) + "\n"
                text += "凶位：" + "｜".join(x) + "\n\n"
                text += f"【四、最佳朝向】\n"
                text += f"大门宜朝：{caifeng}｜床头宜朝：{shiye}\n"
                text += f"书桌宜朝：{wenchang}\n\n"
                text += f"【五、吉利楼层(尾数)】\n"
                text += f"适合：{','.join(map(str, he_lou))} 尾号\n\n"
                text += f"【六、居家布局建议】\n"
                text += f"• 主卧宜静不宜冲\n• 书房文昌位利学业事业\n• 客厅明亮财位整洁\n• 厨房忌对卧室门\n• 卫生间不宜居中\n\n"
                text += f"【七、喜用色彩】\n"
                text += f"幸运色：{se}｜喜用五行：{wx}\n\n"
                text += "🧭 以上为传统文化民俗参考，仅供生活布局使用。"
                return text
            except Exception as e:
                return f"⚠️ 风水生成异常：{str(e)}"


        # ============== 显示风水结果 ==============
        fengshui_txt = get_fengshui_advice()
        st.markdown("---")
        # 读取姓名
        user_name = st.session_state.get("name", "命主")
        st.success(f"✅ {user_name}{gender}｜{shengxiao}｜{bazi_str} 专属风水")
        st.markdown(fengshui_txt)
# ===================== 最终版AI解读+导出Word/PDF =====================
if st.session_state.bottom_nav_active == "解读":
    st.markdown("<div style='text-align:center; margin-top:20px;'><h3>🤖 AI深度命理解读</h3></div>", unsafe_allow_html=True)
    if "bazi_result" not in st.session_state or not st.session_state.bazi_result:
        st.warning("⚠️ 请先在排盘页完成排盘，再使用AI解读")
    else:
        bazi_data = st.session_state.bazi_result
        gender = st.session_state.get("gender", "先生")
        fengshui_txt = ""
        if st.session_state.get("bottom_nav_active") == "风水":
            fengshui_txt = st.session_state.get("fengshui_result", "")
        if st.button("🚀 开始深度解读", use_container_width=True):
            with st.spinner("🤖 AI正在深度分析中...\n⏳ 预计5-15秒，请不要重复点击"):
                try:
                    prompt = f"""你是专业子平八字命理师，按以下结构生成**专业版报告**，严格输出，不要客套话：
【一、命局本质】
八字：{bazi_data['八字_str']}，性别{gender}，日主{bazi_data['日干']}({bazi_data['日干五行']})，
生肖：{bazi_data['生肖']}，五行统计：{bazi_data['五行']}
用生活化比喻总结命局。
【二、寿缘与关键风险年份】
1. 寿缘参考（命理推导，非绝对）:结合八字五行平衡、用神力量、大运走势，推算并明确给出最长寿缘期望参考值（必须是具体年龄范围，如88-95岁）
2. 5个关键风险年份：最近的关键风险年份
3. 每一年注意事项
【三、事业细分】
结合时代判断学历层次、适合行业、岗位、晋升年份
涉及具体的年份要备注（如丙午年（2026）
【四、财运细分】
财富如何、正财偏财、求财方位、风险年份
涉及具体的年份要备注（如丙午年（2026）
【五、婚姻家庭】
配偶情况、相处模式、助力来源、风险年份、预判生儿生女、子女学历、子女职业、子女未来成就如何
涉及具体的年份要备注（如丙午年（2026）
【六、健康细分】
重点养护部位、体检方向
【七、关键节点行动】
凶年：守成、不投资、不跳槽、不远行、一些重大事件(故)何时发生
吉年：主动拓展、合作、考证、一生中有哪些大事件发生
涉及具体的年份要备注（如丙午年（2026）
【八、风水与化解】
吉方位、吉颜色、饰品、布局建议
要求：语言专业、简练、命理师风格，不要符号,禁止出现任何星号（*）和井号（#）、不要标题格式,所有内容基于八字核心数据（空亡、地支互动、用忌神等），不空谈。
"""
                    DEEPSEEK_CONFIG = {
                        "api_key": "sk-d3b976c41272460eab55726001606b15",
                        "api_url": "https://api.deepseek.com/v1/chat/completions",
                        "model": "deepseek-chat",
                        "timeout": 70,
                        "temperature": 0.7,
                        "max_tokens": 2800
                    }
                    headers = {"Content-Type": "application/json","Authorization": f"Bearer {DEEPSEEK_CONFIG['api_key']}"}
                    payload = {"model": DEEPSEEK_CONFIG["model"],"messages": [{"role": "user", "content": prompt}],"temperature": 0.7,"max_tokens": 2800}
                    response = requests.post(DEEPSEEK_CONFIG["api_url"], headers=headers, json=payload, timeout=70)
                    response.raise_for_status()
                    ai_result = response.json()["choices"][0]["message"]["content"].strip()
                    st.session_state.ai_result = ai_result
                    st.markdown("---")
                    st.success("✅ AI深度解读完成")
                    st.markdown(ai_result)
                except Exception as e:
                    st.error(f"❌ 解读失败：{str(e)}")
        # ===================== 【正确版】只在解读页显示导出，不跑顶 =====================
        if "ai_result" in st.session_state and st.session_state.ai_result:
            st.markdown("---")
            st.markdown("#### 📄 导出报告")

            if DOCX_AVAILABLE:
                # 读取姓名
                user_name = st.session_state.get("name", "命主")
                timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                file_name = f"{user_name}_八字解读报告_{timestamp}.docx"

                word_bytes = generate_word_doc(
                    bazi_data,
                    gender,
                    st.session_state.ai_result,
                    st.session_state.get("fengshui_result", "")
                )

                st.download_button(
                    label="📄 导出Word",
                    data=word_bytes,
                    file_name=file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="download_word_report"  # 解决重复ID报错
                )
            else:
                st.button("📄 导出Word（未安装python-docx）", disabled=True, use_container_width=True)
# ========= 【底部固定悬浮导航栏】只负责跟随状态变色，纯视觉层 =========
st.markdown("<div style='height:90px;'></div>",unsafe_allow_html=True)
act = st.session_state.bottom_nav_active
# 判定高亮样式
s1 = "color:#D4AF37;font-weight:bold;" if act=="排盘" else "color:#555;"
s2 = "color:#D4AF37;font-weight:bold;" if act=="吉日" else "color:#555;"
s3 = "color:#D4AF37;font-weight:bold;" if act=="风水" else "color:#555;"
s4 = "color:#D4AF37;font-weight:bold;" if act=="解读" else "color:#555;"
s5 = "color:#D4AF37;font-weight:bold;" if act=="." else "color:#555;"
s6 = "color:#D4AF37;font-weight:bold;" if act==".." else "color:#555;"
s7 = "color:#D4AF37;font-weight:bold;" if act=="..." else "color:#555;"
# 底部固定7项导航（永久悬浮底部）
st.markdown(f"""
<style>
.bottom-fixed-nav{{
    position:fixed;bottom:0;left:0;right:0;
    background:#fff;border-top:1px solid #eee;
    display:flex;padding:10px 0;z-index:9999;
}}
.nav-item{{flex:1;text-align:center;font-size:14px;cursor:default;}}
</style>
<div class="bottom-fixed-nav">
    <div class="nav-item" style="{s1}">☯️<br>排盘</div>
    <div class="nav-item" style="{s2}">⏱️<br>吉日</div>
    <div class="nav-item" style="{s3}">🏮<br>风水</div>
    <div class="nav-item" style="{s4}">📄<br>解读</div>
    <div class="nav-item" style="{s5}"></div>
    <div class="nav-item" style="{s6}"></div>
    <div class="nav-item" style="{s7}"></div>
</div>
""",unsafe_allow_html=True)


